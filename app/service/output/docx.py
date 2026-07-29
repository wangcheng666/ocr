# coding=utf-8
"""将 MinerU OCR 中间结果 middle.json 转换成 docx 文件。

要求：
- 内容不要丢失，顺序排列
- 页面和 pdf 保持对应
- 格式按照公文格式
"""

import io
from collections import defaultdict
from enum import Enum

import latex2mathml.converter
import mathml2omml
from bs4 import BeautifulSoup
from docx import Document
from docx.document import Document as DocumemtObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell
from loguru import logger
from lxml import etree
from mineru.data.data_reader_writer.base import DataReader, DataWriter
from mineru.utils.enum_class import BlockType, ContentType

# ──────────────────────────────────────────────
# 命名空间常量
# ──────────────────────────────────────────────
NS_M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ──────────────────────────────────────────────
# 公文样式定义
# ──────────────────────────────────────────────
class DocContentType(Enum):
    """文档内容类型"""
    TITLE = 1
    FIRST_TITLE = 2
    SECOND_TITLE = 3
    THIRD_TITLE = 4
    FOURTH_TITLE = 5
    MAIN_BODY = 6


_STYLES = {
    DocContentType.TITLE: {
        "alignment": "center",
        "first_line_indent": 0,
        "font": {"size": 22, "name": "方正小标宋简体", "bold": False, "italic": False},
    },
    DocContentType.FIRST_TITLE: {
        "alignment": "center",
        "first_line_indent": 0,
        "font": {"size": 16, "name": "黑体", "bold": False, "italic": False},
    },
    DocContentType.SECOND_TITLE: {
        "alignment": "left",
        "first_line_indent": 0,
        "font": {"size": 16, "name": "楷体", "bold": False, "italic": False},
    },
    DocContentType.THIRD_TITLE: {
        "alignment": "left",
        "first_line_indent": 0,
        "font": {"size": 16, "name": "仿宋", "bold": False, "italic": False},
    },
    DocContentType.FOURTH_TITLE: {
        "alignment": "left",
        "first_line_indent": 0,
        "font": {"size": 16, "name": "仿宋", "bold": False, "italic": False},
    },
    DocContentType.MAIN_BODY: {
        "alignment": "left",
        "first_line_indent": 0.42,
        "font": {"size": 16, "name": "仿宋", "bold": False, "italic": False},
    },
}
_level_type_map = {m.value: m for m in DocContentType}


def _set_style(paragraph, style_type: DocContentType):
    """对段落应用公文样式。"""
    s = _STYLES[style_type]
    if s["alignment"] == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif s["alignment"] == "left":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Inches(s["first_line_indent"])
    if not paragraph.runs:
        return
    f = s["font"]
    for run in paragraph.runs:
        run.font.size = Pt(f["size"])
        run.bold = f["bold"]
        run.italic = f["italic"]
        run.font.name = f["name"]
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), f["name"])


# ──────────────────────────────────────────────
# 公式处理：LaTeX → MathML → OMML
# ──────────────────────────────────────────────
def _latex_to_mathml(latex: str) -> str:
    """LaTeX → MathML XML 字符串。"""
    try:
        result = latex2mathml.converter.convert(latex)
        elem = etree.fromstring(result.encode("utf-8")) if isinstance(result, str) else result
        return etree.tostring(elem, pretty_print=True, encoding="unicode")
    except Exception as e:
        logger.error(f"LaTeX→MathML 失败: {e}")
        return ""


def _insert_omml_formula(paragraph, omml_xml: str):
    """将 OMML <m:oMath> 注入到 Word 段落中。"""
    if "xmlns:m" not in omml_xml:
        omml_xml = omml_xml.replace(
            "<m:oMath>", f'<m:oMath xmlns:m="{NS_M}" xmlns:w="{NS_W}">', 1
        )
        if "xmlns:m" not in omml_xml:
            raise ValueError("无法找到 <m:oMath> 标签")
    try:
        run_el = OxmlElement("w:r")
        parser = etree.XMLParser(recover=True)
        run_el.append(etree.fromstring(omml_xml.encode("utf-8"), parser))
        paragraph._element.append(run_el)
    except Exception as e:
        logger.error(f"注入 OMML 公式失败: {e}")
        paragraph.add_run("[公式注入失败]")


# ──────────────────────────────────────────────
# 表格 HTML 解析 & 写入
# ──────────────────────────────────────────────
def _split_rectangular_tables(rows):
    """按列数变化将不规则表格拆分成多个矩形子表格。"""
    if not rows:
        return []
    tables, cur, cur_cols = [], [], len(rows[0])
    for row in rows:
        if len(row) == cur_cols:
            cur.append(row)
        else:
            if cur:
                tables.append(cur)
            cur, cur_cols = [row], len(row)
    if cur:
        tables.append(cur)
    return tables


def _parse_html_table(html: str):
    """解析 HTML 表格（含 colspan/rowspan）为矩形列表矩阵。"""
    soup = BeautifulSoup(html, "lxml")
    table = soup.find("table")
    if not table:
        return []
    occupied = defaultdict(int)
    grid = []
    for tr in table.find_all("tr"):
        row, col = [], 0
        for cell in tr.find_all(["td", "th"]):
            text = cell.get_text(strip=True)
            while col in occupied:
                row.append("v")
                occupied[col] -= 1
                if occupied[col] == 0:
                    del occupied[col]
                col += 1
            cs = int(cell.get("colspan", 1))
            rs = int(cell.get("rowspan", 1))
            for i in range(cs):
                row.append(text if i == 0 else "h")
            if rs > 1:
                for c in range(col, col + cs):
                    occupied[c] = rs - 1
            col += cs
        while col in occupied:
            row.append("v")
            occupied[col] -= 1
            if occupied[col] == 0:
                del occupied[col]
            col += 1
        grid.append(row)
    return _split_rectangular_tables(grid)


def _write_table(doc, grid):
    """将矩形表格（含合并标记）写入 docx。"""
    rows, cols = len(grid), len(grid[0]) if grid else 0
    if rows == 0:
        return
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    v_merges = {}
    for r, row in enumerate(grid):
        doc_row = table.rows[r]
        h_start = None
        for c, val in enumerate(row):
            cell = doc_row.cells[c]
            if h_start:
                if val == "h":
                    continue
                h_start.merge(doc_row.cells[c - 1])
                h_start = None
            if c in v_merges:
                if val == "v":
                    try:
                        v_merges[c][1].merge(cell)
                        if r == rows - 1:
                            del v_merges[c]
                    except Exception as e:
                        logger.error(f"垂直合并失败 R{r}C{c}: {e}")
                    continue
                del v_merges[c]
            cell.text = val
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if r == 0:
                    cell.paragraphs[0].runs[0].bold = True
            if c + 1 < cols and row[c + 1] == "h":
                h_start = cell
            if r + 1 < rows and grid[r + 1][c] == "v":
                v_merges[c] = (r, cell)
        if h_start:
            h_start.merge(doc_row.cells[cols - 1])


# ══════════════════════════════════════════════
# 内容块写入器 — DocxGenerator 类
# ══════════════════════════════════════════════

def _write_centered_text(doc, block, *, italic: bool):
    """写入居中文本（图/表标题或脚注）。"""
    text = "".join(
        span["content"]
        for line in block.get("lines", [])
        for span in line.get("spans", [])
        if span["type"] == ContentType.TEXT
    )
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    if italic:
        run.italic = True


class DocxGenerator:
    """DOCX 生成器，持有全局上下文（img_reader 等），避免逐层传递。"""

    def __init__(
        self,
        img_reader: DataReader,
        *,
        formula_enable: bool = True,
        process_formula: bool = True,
        table_enable: bool = True,
    ):
        self.img_reader = img_reader
        self.formula_enable = formula_enable
        self.process_formula = process_formula
        self.table_enable = table_enable

    # ── 文本 ──
    def _write_text(
        self, doc: DocumemtObject, block, new_paragraph: bool = True
    ):
        """写入文本段落（TEXT / INLINE_EQUATION / INTERLINE_EQUATION）。"""
        para = doc.add_paragraph() if new_paragraph else doc.paragraphs[-1]
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                st = span["type"]
                content = span["content"]
                if st == ContentType.TEXT:
                    para.add_run(content)
                elif st == ContentType.INLINE_EQUATION:
                    if self.process_formula:
                        _insert_omml_formula(
                            para, mathml2omml.convert(_latex_to_mathml(content))
                        )
                    else:
                        para.add_run(content)
                elif st == ContentType.INTERLINE_EQUATION:
                    if self.formula_enable:
                        if self.process_formula:
                            _set_style(para, DocContentType.MAIN_BODY)
                            para = doc.add_paragraph()
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            _insert_omml_formula(
                                para, mathml2omml.convert(_latex_to_mathml(content))
                            )
                        else:
                            para.add_run(content)
                    else:
                        img_path = span.get("image_path", "")
                        if img_path:
                            self._write_img(doc, img_path)
        _set_style(para, DocContentType.MAIN_BODY)

    # ── 标题 ──
    def _write_title(self, doc, block, level: int):
        text = "".join(
            span["content"]
            for line in block.get("lines", [])
            for span in line.get("spans", [])
            if span["type"] == ContentType.TEXT
        )
        style = _level_type_map.get(level, DocContentType.FOURTH_TITLE)
        para = doc.add_heading(text=text, level=level)
        _set_style(para, style)

    # ── 图片 ──
    def _write_img(self, doc, url: str):
        """向文档插入单张图片。"""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            data = self.img_reader.read_at(url)
            para.add_run().add_picture(io.BytesIO(data), width=Inches(6.0))
        except Exception as e:
            para.add_run(f"[图片失败: {url}]")
            logger.error(f"图片插入失败 {url}: {e}")

    def _write_image_block(self, doc, block):
        """写入图片复合块（body / caption / footnote）。"""
        has_fn = any(
            b["type"] == BlockType.IMAGE_FOOTNOTE for b in block.get("blocks", [])
        )
        order = (
            [BlockType.IMAGE_CAPTION, BlockType.IMAGE_BODY, BlockType.IMAGE_FOOTNOTE]
            if has_fn
            else [BlockType.IMAGE_BODY, BlockType.IMAGE_CAPTION]
        )
        for target in order:
            for sub in block.get("blocks", []):
                if sub["type"] != target:
                    continue
                if target == BlockType.IMAGE_BODY:
                    self._write_image_body(doc, sub)
                else:
                    _write_centered_text(
                        doc, sub, italic=(target != BlockType.IMAGE_FOOTNOTE)
                    )

    def _write_image_body(self, doc, block):
        """写入图片数据体。"""
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                if span["type"] == ContentType.IMAGE:
                    path = span.get("image_path", "")
                    if path:
                        self._write_img(doc, path)

    # ── 表格 ──
    def _write_table_block(self, doc, block):
        """写入表格复合块（caption / body / footnote）。"""
        for sub in block.get("blocks", []):
            t = sub["type"]
            if t == BlockType.TABLE_CAPTION:
                _write_centered_text(doc, sub, italic=True)
            elif t == BlockType.TABLE_BODY:
                self._write_table_data(doc, sub)
            elif t == BlockType.TABLE_FOOTNOTE:
                _write_centered_text(doc, sub, italic=False)

    def _write_table_data(self, doc, block):
        """写入表格数据体。"""
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                if span["type"] != ContentType.TABLE:
                    continue
                if self.table_enable:
                    html = span.get("html", "")
                    if html:
                        try:
                            for grid in _parse_html_table(html):
                                _write_table(doc, grid)
                            return
                        except Exception as e:
                            logger.error(f"表格解析失败: {e}")
                path = span.get("image_path", "")
                if path:
                    self._write_img(doc, path)

    # ── 列表 ──
    def _write_list(self, doc, list_block):
        """递归写入列表块（支持嵌套子列表）。"""
        for block in list_block.get("blocks", []):
            if block.get("type") == BlockType.LIST:
                self._write_list(doc, block)
            else:
                self._write_text(doc, block)

    # ── 代码 ──
    def _write_code(self, doc, block):
        """写入代码复合块。"""
        for sub in block.get("blocks", []):
            if sub["type"] in (BlockType.CODE_CAPTION, BlockType.CODE_BODY):
                self._write_text(doc, sub)

    # ── 主入口 ──
    def generate(self, pdf_info_dict) -> DocumemtObject:
        """将中间结果 pdf_info_dict 转换为 python-docx Document。"""
        document = Document()
        for page_info in pdf_info_dict:
            blocks = page_info.get("para_blocks", [])
            new_para = not _is_paragraph_cross_page(document)
            for i, block in enumerate(blocks):
                t = block["type"]
                if t in (
                    BlockType.TEXT,
                    BlockType.INTERLINE_EQUATION,
                    BlockType.PHONETIC,
                    BlockType.REF_TEXT,
                ):
                    self._write_text(document, block, new_para)
                elif t == BlockType.LIST:
                    self._write_list(document, block)
                elif t == BlockType.TITLE:
                    self._write_title(document, block, _get_title_level(block))
                elif t == BlockType.IMAGE:
                    self._write_image_block(document, block)
                elif t == BlockType.TABLE:
                    self._write_table_block(document, block)
                elif t == BlockType.CODE:
                    self._write_code(document, block)
                if i == 0:
                    new_para = True
        return document


# ──────────────────────────────────────────────
# 辅助函数
# ──────────────────────────────────────────────
_PUNCTUATION = set("，。！？；：“”【】（）《》、,.!?;:\"'()[]{}")


def _is_paragraph_cross_page(doc) -> bool:
    """根据末尾标点判断上一段落是否跨页（未结束）。"""
    if not doc.paragraphs:
        return False
    last = doc.paragraphs[-1]
    text = last.text.strip()
    if last.style.name.startswith("Heading"):
        return False
    return bool(text) and text[-1] not in _PUNCTUATION


def _get_title_level(block) -> int:
    """获取标题等级（python-docx 支持 1-9）。"""
    return max(block.get("level", 1), 0)


def to_docx(
    img_reader: DataReader,
    file_writer: DataWriter,
    file_name: str,
    pdf_info_dict,
    formula_enable: bool = True,
    table_enable: bool = True,
    f_dump_docx: bool = True,
):
    """将中间结果转换为 docx 并写入 S3。"""
    if not f_dump_docx:
        logger.info("f_dump_docx=False，跳过 docx 生成")
        return
    for suffix, pf in [("", True), ("_with_raw_formula", False)]:
        gen = DocxGenerator(
            img_reader=img_reader,
            formula_enable=formula_enable,
            process_formula=pf,
            table_enable=table_enable,
        )
        doc = gen.generate(pdf_info_dict)
        buf = io.BytesIO()
        doc.save(buf)
        file_writer.write(f"{file_name}{suffix}.docx", data=buf.getvalue())
