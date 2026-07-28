#!/usr/bin/env python3
"""创建用于前后对比测试的 DOCX 文档，覆盖多种编号格式场景。"""

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import copy
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "numbering_test.docx")

# ── 命名空间 ──────────────────────────────────────────────
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
nsmap = {"w": W, "r": R}


def _make_val_attr(name: str, val: str) -> dict:
    return {f"{{{W}}}{name}": val}


def make_numbering_xml():
    """构造 numbering.xml 内容，包含多种 numFmt 的编号定义。"""
    NUM_FMTS = [
        # (abstractNumId, ilvl0_numFmt, ilvl0_lvlText, ilvl1_numFmt, ilvl1_lvlText)
        (0, "decimal", "%1.", "lowerLetter", "%2)"),       # 1. / a)
        (1, "chineseCounting", "%1、", "decimal", "(%2)"),  # 一、 / (1)
        (2, "lowerRoman", "%1.", "decimal", "%1.%2"),      # i. / 1.1
        (3, "bullet", "●", "bullet", "○"),                  # bullet nested
        (4, "upperRoman", "%1.", "decimal", "%1.%2"),      # I. / 1.1
        (5, "decimalEnclosedCircle", "(%1)", None, None),  # ①
        (6, "chineseCountingThousand", "%1、", None, None), # 壹、贰、
    ]

    xml = '<w:numbering %s>\n' % nsdecls("w")
    xml += '  <w:numIdMacAtCleanUp w:val="0"/>\n'

    for abs_id, fmt0, text0, fmt1, text1 in NUM_FMTS:
        # abstractNum
        xml += f'  <w:abstractNum w:abstractNumId="{abs_id}">\n'
        xml += '    <w:nsid w:val="FFFFFFFF"/>\n'
        xml += '    <w:multiLevelType w:val="multilevel"/>\n'
        xml += '    <w:tmpl w:val="FFFFFFFF"/>\n'
        # level 0
        xml += f'    <w:lvl w:ilvl="0">\n'
        xml += f'      <w:start w:val="1"/>\n'
        xml += f'      <w:numFmt w:val="{fmt0}"/>\n'
        xml += f'      <w:lvlText w:val="{text0}"/>\n'
        xml += '      <w:lvlJc w:val="left"/>\n'
        xml += '      <w:pPr>\n'
        xml += '        <w:ind w:left="360" w:hanging="360"/>\n'
        xml += '      </w:pPr>\n'
        xml += '    </w:lvl>\n'

        if fmt1 is not None:
            xml += f'    <w:lvl w:ilvl="1">\n'
            xml += f'      <w:start w:val="1"/>\n'
            xml += f'      <w:numFmt w:val="{fmt1}"/>\n'
            xml += f'      <w:lvlText w:val="{text1}"/>\n'
            xml += '      <w:lvlJc w:val="left"/>\n'
            xml += '      <w:pPr>\n'
            xml += '        <w:ind w:left="720" w:hanging="360"/>\n'
            xml += '      </w:pPr>\n'
            xml += '    </w:lvl>\n'

        xml += '  </w:abstractNum>\n'

    # num instances (each numId = abstractNumId + 1)
    for abs_id, *_ in NUM_FMTS:
        num_id = abs_id + 1
        xml += f'  <w:num w:numId="{num_id}">\n'
        xml += f'    <w:abstractNumId w:val="{abs_id}"/>\n'
        xml += '  </w:num>\n'

    xml += '</w:numbering>'
    return xml


def add_numbered_paragraph(doc, text: str, numId: int, ilvl: int = 0):
    """添加带编号的段落。"""
    p = doc.add_paragraph(text)
    pPr = p._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p._element.insert(0, pPr)
    numPr = parse_xml(
        f'<w:numPr {nsdecls("w")}>'
        f'  <w:ilvl w:val="{ilvl}"/>'
        f'  <w:numId w:val="{numId}"/>'
        f'</w:numPr>'
    )
    pPr.append(numPr)
    return p


def main():
    doc = Document()

    # ── 替换 numbering.xml ───────────────────────────────
    numbering_xml_str = make_numbering_xml()
    for part in doc.part.package.parts:
        pn = part.partname.lower() if hasattr(part, "partname") else str(getattr(part, "partname", "")).lower()
        if "numbering" in pn:
            root = etree.fromstring(numbering_xml_str.encode())
            part._element = root
            break

    # ── 写标题 ───────────────────────────────────────────
    doc.add_heading("DOCX 编号格式对比测试文档", level=0)

    # ── 场景 1: decimal 有序列表 ─────────────────────────
    doc.add_heading("1. decimal 有序列表 (numId=1)", level=1)
    add_numbered_paragraph(doc, "第一项（应有编号 1.）", numId=1, ilvl=0)
    add_numbered_paragraph(doc, "第二项（应有编号 2.）", numId=1, ilvl=0)
    doc.add_paragraph("（这里插入一段正文，用于测试列表中断）")
    add_numbered_paragraph(doc, "第三项（中断后继续，应有编号 3.）", numId=1, ilvl=0)

    # ── 场景 2: 中文编号 (chineseCounting) ───────────────
    doc.add_heading("2. chineseCounting 中文编号 (numId=2)", level=1)
    add_numbered_paragraph(doc, "必要性分析（应有编号 一、）", numId=2, ilvl=0)
    add_numbered_paragraph(doc, "研究现状（应有编号 二、）", numId=2, ilvl=0)
    add_numbered_paragraph(doc, "    子项 a（应有编号 (1)）", numId=2, ilvl=1)
    add_numbered_paragraph(doc, "    子项 b（应有编号 (2)）", numId=2, ilvl=1)
    add_numbered_paragraph(doc, "实验设计（应有编号 三、）", numId=2, ilvl=0)

    # ── 场景 3: 罗马数字 (lowerRoman) ────────────────────
    doc.add_heading("3. lowerRoman 罗马数字 (numId=3)", level=1)
    add_numbered_paragraph(doc, "Introduction（应有编号 i.）", numId=3, ilvl=0)
    add_numbered_paragraph(doc, "Background（应有编号 ii.）", numId=3, ilvl=0)
    add_numbered_paragraph(doc, "    子节 1（应有编号 1.1）", numId=3, ilvl=1)
    add_numbered_paragraph(doc, "    子节 2（应有编号 1.2）", numId=3, ilvl=1)
    add_numbered_paragraph(doc, "Conclusion（应有编号 iii.）", numId=3, ilvl=0)

    # ── 场景 4: bullet 无序列表 ──────────────────────────
    doc.add_heading("4. bullet 无序列表 (numId=4)", level=1)
    add_numbered_paragraph(doc, "苹果（应有 ●）", numId=4, ilvl=0)
    add_numbered_paragraph(doc, "香蕉（应有 ●）", numId=4, ilvl=0)
    add_numbered_paragraph(doc, "    热带水果（应有 ○）", numId=4, ilvl=1)
    add_numbered_paragraph(doc, "    浆果类（应有 ○）", numId=4, ilvl=1)
    add_numbered_paragraph(doc, "橘子（应有 ●）", numId=4, ilvl=0)

    # ── 场景 5: 大写罗马数字 (upperRoman) ────────────────
    doc.add_heading("5. upperRoman 大写罗马数字 (numId=5)", level=1)
    add_numbered_paragraph(doc, "Part One（应有编号 I.）", numId=5, ilvl=0)
    add_numbered_paragraph(doc, "Part Two（应有编号 II.）", numId=5, ilvl=0)

    # ── 场景 6: heading-style 列表（带正文穿插）─────────
    doc.add_heading("6. heading-style 列表（正文穿插，应转 title 块）", level=1)
    add_numbered_paragraph(doc, "第一章 引言（应转为 title 且前缀为 一、）", numId=2, ilvl=0)
    doc.add_paragraph("这是第一章下面的正文内容……正文穿插是 heading 列表的判断条件之一。")
    add_numbered_paragraph(doc, "第二章 方法（应转为 title 且前缀为 二、）", numId=2, ilvl=0)
    doc.add_paragraph("这是第二章下面的正文内容……")
    add_numbered_paragraph(doc, "第一节 数据采集（子标题应有 (1)）", numId=2, ilvl=1)
    doc.add_paragraph("数据采集的方法描述……")
    add_numbered_paragraph(doc, "第二节 数据分析（子标题应有 (2)）", numId=2, ilvl=1)
    doc.add_paragraph("数据分析的详细说明……")
    add_numbered_paragraph(doc, "第三章 结果（应转为 title 且前缀为 三、）", numId=2, ilvl=0)

    # ── 场景 7: 表格中断列表 ─────────────────────────────
    doc.add_heading("7. 表格中断后的列表（应重新编号）", level=1)
    add_numbered_paragraph(doc, "列表项 A（应有编号 1.）", numId=1, ilvl=0)
    add_numbered_paragraph(doc, "列表项 B（应有编号 2.）", numId=1, ilvl=0)
    doc.add_table(rows=2, cols=2)
    # 表格后 numId=1 的计数器被 reset，应重新从 1 开始
    add_numbered_paragraph(doc, "表格后列表项（应有编号 1.）", numId=1, ilvl=0)
    add_numbered_paragraph(doc, "表格后列表项（应有编号 2.）", numId=1, ilvl=0)

    # ── 场景 8: 多层嵌套列表 ─────────────────────────────
    doc.add_heading("8. 多层嵌套列表 (numId=26-like)", level=1)
    add_numbered_paragraph(doc, "一级项 1", numId=1, ilvl=0)
    add_numbered_paragraph(doc, "    二级项 a", numId=1, ilvl=1)
    add_numbered_paragraph(doc, "        三级项 i", numId=1, ilvl=2)
    add_numbered_paragraph(doc, "    二级项 b", numId=1, ilvl=1)
    add_numbered_paragraph(doc, "一级项 2", numId=1, ilvl=0)

    # ── 保存 ─────────────────────────────────────────────
    doc.save(OUTPUT)
    print(f"测试文档已创建: {OUTPUT}")
    print(f"场景覆盖:")
    print(f"  1. decimal 有序列表 (numId=1)")
    print(f"  2. chineseCounting 中文编号 (numId=2)")
    print(f"  3. lowerRoman 罗马数字 (numId=3)")
    print(f"  4. bullet 无序列表 (numId=4)")
    print(f"  5. upperRoman 大写罗马数字 (numId=5)")
    print(f"  6. heading-style 列表（正文穿插）")
    print(f"  7. 表格中断后重新编号")
    print(f"  8. 多层嵌套列表")


if __name__ == "__main__":
    main()
